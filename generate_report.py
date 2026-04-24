
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    def add_title_page():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("A PROJECT REPORT ON\n\n")
        run.bold = True
        run.size = Pt(16)
        
        run = p.add_run("HYBRID AI INTELLIGENCE DASHBOARD FOR FRAMING BIAS, FAKE NEWS, AND MANIPULATION DETECTION\n\n\n")
        run.bold = True
        run.size = Pt(20)
        
        run = p.add_run("BACHELOR OF TECHNOLOGY\n")
        run.size = Pt(14)
        run = p.add_run("IN\n")
        run = p.add_run("COMPUTER SCIENCE AND INFORMATION TECHNOLOGY (AIML / CYBER SECURITY)\n\n\n")
        run.bold = True
        
        run = p.add_run("SUBMITTED BY\n")
        run = p.add_run("GROUP NO: AI-07\n\n")
        
        table = doc.add_table(rows=3, cols=2)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(0,0).text = "PRN"
        table.cell(0,1).text = "NAME OF THE STUDENT"
        table.cell(1,0).text = "12345678"
        table.cell(1,1).text = "Animesh Pathak"
        table.cell(2,0).text = "87654321"
        table.cell(2,1).text = "Team Partner"

        doc.add_paragraph("\n\n\n")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("UNDER THE GUIDANCE OF\n")
        run = p.add_run("PROF. [GUIDE NAME]\n\n\n")
        
        run = p.add_run("SCHOOL OF COMPUTER SCIENCE AND INFORMATION TECHNOLOGY\n")
        run = p.add_run("SYMBIOSIS SKILLS AND PROFESSIONAL UNIVERSITY, PUNE\n")
        run = p.add_run("ACADEMIC YEAR 2025-26")
        doc.add_page_break()

    def add_toc_lists():
        doc.add_heading('TABLE OF CONTENTS', 0)
        doc.add_paragraph("Abstract")
        doc.add_paragraph("List of Figures")
        doc.add_paragraph("List of Tables")
        doc.add_paragraph("List of Abbreviations")
        doc.add_paragraph("Chapter 1: Introduction")
        doc.add_paragraph("Chapter 2: Literature Review")
        doc.add_paragraph("Chapter 3: Subject-wise Integration")
        doc.add_paragraph("Chapter 4: System Design")
        doc.add_paragraph("Chapter 5: Implementation")
        doc.add_paragraph("Chapter 6: Results and Analysis")
        doc.add_paragraph("Chapter 7: Discussion")
        doc.add_paragraph("Chapter 8: Conclusion & Future Scope")
        doc.add_page_break()

        doc.add_heading('LIST OF ABBREVIATIONS', 0)
        abbrs = [
            ["NLP", "Natural Language Processing"],
            ["AI", "Artificial Intelligence"],
            ["ML", "Machine Learning"],
            ["LLM", "Large Language Model"],
            ["API", "Application Programming Interface"],
            ["TF-IDF", "Term Frequency-Inverse Document Frequency"],
            ["JSON", "JavaScript Object Notation"],
            ["REST", "Representational State Transfer"],
            ["LIME", "Local Interpretable Model-agnostic Explanations"],
            ["CORS", "Cross-Origin Resource Sharing"],
            ["UI/UX", "User Interface / User Experience"]
        ]
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Abbreviation'
        hdr[1].text = 'Expansion'
        for a, e in abbrs:
            row = table.add_row().cells
            row[0].text = a
            row[1].text = e
        doc.add_page_break()

    def add_abstract():
        doc.add_heading('ABSTRACT', 0)
        text = (
            "In the modern digital era, the proliferation of misinformation and framing bias has become a significant "
            "threat to democratic discourse and public security. Traditional fact-checking methods are often slow and "
            "struggle to keep pace with the velocity of social media. This project proposes a 'Hybrid AI Intelligence Dashboard' "
            "that integrates classical Machine Learning (ML) for high-speed classification, Natural Language Processing (NLP) "
            "for linguistic threat analysis, and Large Language Models (LLMs) for deep contextual reasoning. "
            "The system utilizes a 4-stage pipeline: (1) ML-based framing detection, (2) Lexical threat scoring, "
            "(3) Automated live news context retrieval via DuckDuckGo, and (4) LLM-driven comparative analysis using LLaMA 3.3. "
            "Results indicate a classification accuracy of 83.2% on diverse news datasets, with the system providing "
            "real-time transparency through LIME-based explainability and security logs for auditability. "
            "This project bridges the gap between static analysis and dynamic real-world verification, offering a robust "
            "tool for identifying subtle media manipulation."
        )
        doc.add_paragraph(text)
        doc.add_page_break()

    def add_introduction():
        doc.add_heading('CHAPTER 1: INTRODUCTION', 1)
        
        doc.add_heading('1.1 Background', 2)
        doc.add_paragraph(
            "The digital landscape has transformed news consumption, shifting it from traditional broadcast media to "
            "decentralized social platforms. While this democratizes information access, it also facilitates the rapid spread "
            "of 'Fake News' and 'Framing Bias'. Framing bias occurs when information is presented in a way that encourages "
            "a specific interpretation, often by highlighting certain facts while omitting others. Detecting these "
            "subtle linguistic nuances requires more than simple keyword matching; it demands a multi-layered computational "
            "approach that combines statistical modeling with deep semantic understanding."
        )

        doc.add_heading('1.2 Motivation', 2)
        doc.add_paragraph(
            "The motivation behind this project stems from the increasing impact of 'Information Warfare' and 'Cognitive Hacking'. "
            "Misinformation is no longer just a nuisance; it is a cybersecurity threat that can manipulate elections, "
            "incite violence, and destabilize financial markets. By creating a tool that can decode these hidden agendas "
            "in real-time, we aim to empower users and analysts with 'Digital Resilience'. The ability to verify a "
            "breaking news story against a global context within seconds provides a critical defense against psychological manipulation."
        )

        doc.add_heading('1.3 Problem Statement', 2)
        doc.add_paragraph(
            "Current automated fact-checking systems often suffer from two major flaws: (1) High Latency: Deep analysis "
            "is computationally expensive, and (2) Lack of Context: Systems analyze text in isolation without considering "
            "live events. There is a critical need for a hybrid architecture that leverages the speed of local ML models "
            "and the reasoning power of cloud-based LLMs, while simultaneously fetching live context to validate claims."
        )

        doc.add_heading('1.4 Objectives', 2)
        objs = [
            "To design and implement a hybrid 4-stage AI pipeline for news analysis.",
            "To integrate local ML models for real-time framing classification with >80% accuracy.",
            "To implement lexical threat scoring for detecting propaganda and emotional manipulation.",
            "To develop an automated context-fetching service using search engine APIs.",
            "To utilize LLaMA 3.3 (via GroqCloud) for generating deep comparative reports and human-readable verdicts.",
            "To provide a secure, interactive dashboard for visualization and security logging."
        ]
        for obj in objs:
            doc.add_paragraph(obj, style='List Bullet')

        doc.add_heading('1.5 Scope', 2)
        doc.add_paragraph(
            "The scope of this project encompasses the development of a web-based intelligence dashboard. It targets "
            "journalists, cybersecurity analysts, and the general public. The system is designed to handle English-language "
            "news snippets and provides analysis on bias, factuality, and emotional tone. Future iterations could expand "
            "to multi-lingual support and deep-fake video analysis."
        )
        doc.add_page_break()

    def add_literature_review():
        doc.add_heading('CHAPTER 2: LITERATURE REVIEW', 1)
        doc.add_paragraph(
            "A review of existing literature reveals a transition from traditional rule-based systems to deep learning "
            "architectures. Early studies focused on TF-IDF and Naive Bayes for spam detection, which laid the foundation "
            "for modern fake news detection. However, these models fail to capture the 'framing' or 'intent' behind a message."
        )
        
        doc.add_heading('2.1 Existing Solutions', 2)
        doc.add_paragraph(
            "Tools like Snopes and PolitiFact provide high-quality human verification but lack scalability. "
            "On the automated side, systems like 'FakeNewsNet' provide large datasets but are often static. "
            "Our solution differs by introducing a 'Live Context' layer that cross-references inputs with real-time news."
        )

        doc.add_heading('2.2 Comparative Study', 2)
        table = doc.add_table(rows=4, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Feature'
        hdr_cells[1].text = 'Traditional ML'
        hdr_cells[2].text = 'Hybrid AI (Our Project)'
        
        data = [
            ['Speed', 'Fast (Local)', 'Optimized (Local + Async Cloud)'],
            ['Reasoning', 'None (Statistical)', 'Deep (Semantic LLM)'],
            ['Context', 'Static Training Data', 'Live Web Retrieval']
        ]
        for f, t, h in data:
            row = table.add_row().cells
            row[0].text = f
            row[1].text = t
            row[2].text = h
        doc.add_page_break()

    def add_subject_integration():
        doc.add_heading('CHAPTER 3: SUBJECT-WISE INTEGRATION', 1)
        doc.add_paragraph("The project is designed as a multidisciplinary capstone, integrating six core domains of CSIT:")
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Course Code'
        hdr[1].text = 'Course Name'
        hdr[2].text = 'Concept Used'
        hdr[3].text = 'Application in Project'
        
        subjects = [
            ['AI0403', 'NLP', 'TF-IDF, Text Processing', 'ML classification + LLM framing analysis'],
            ['AI0402', 'Web App Dev', 'REST API, Client-Server', 'Flask backend + interactive dashboard'],
            ['AI0404', 'Computer Networks', 'HTTP/HTTPS, API Calls', 'Secure Groq & DuckDuckGo integration'],
            ['AI0401', 'Big Data Analytics', 'Stream Processing', 'Live news parsing & unstructured data'],
            ['AI0405', 'DevOps', 'Modular Pipeline, Env Vars', 'Structured AI pipeline + fallback system'],
            ['AI0406B', 'AI in Cybersecurity', 'Threat Scoring', 'Manipulation detection + security logging']
        ]
        for s in subjects:
            row = table.add_row().cells
            for i in range(4):
                row[i].text = s[i]
        doc.add_page_break()

    def add_system_design():
        doc.add_heading('CHAPTER 4: SYSTEM DESIGN', 1)
        
        doc.add_heading('4.1 Architecture Diagram', 2)
        doc.add_paragraph(
            "The system follows a 'Modular Service-Oriented Architecture' (SOA). The architecture is divided into "
            "three main layers: Client Layer (HTML/JS), Application Layer (Flask API + Pipelines), and External "
            "Service Layer (Groq LLM + Search Engines)."
        )
        
        # Describing a Block Diagram
        doc.add_paragraph("[SYSTEM ARCHITECTURE BLOCK DIAGRAM DESCRIPTION]")
        doc.add_paragraph(
            "1. User Interface (Input Text/URL) -> \n"
            "2. Flask API (app.py) -> \n"
            "3. Pipeline Orchestrator (pipeline.py) -> \n"
            "    a. Local ML Service (Scikit-Learn)\n"
            "    b. NLP Service (Lexical Scorer)\n"
            "    c. Async News Fetcher (DuckDuckGo)\n"
            "    d. Async LLM Service (Llama 3.3)\n"
            "4. Results Merger -> Dashboard Display", style='Quote'
        )

        doc.add_heading('4.2 Hardware Requirement', 2)
        doc.add_paragraph("- CPU: Intel i5 / AMD Ryzen 5 or higher\n- RAM: 8GB minimum (16GB recommended)\n- Storage: 1GB free space\n- Connectivity: High-speed Internet for API requests")

        doc.add_heading('4.3 Software Requirement', 2)
        doc.add_paragraph("- Operating System: Windows 10/11 / Linux / macOS\n- Language: Python 3.11+\n- Framework: Flask 3.0\n- Libraries: Scikit-learn, Pandas, Groq, DuckDuckGo-Search, python-dotenv\n- Browser: Chrome/Firefox (Latest)")

        doc.add_heading('4.5 Database Design', 2)
        doc.add_paragraph(
            "The system utilizes a 'Schema-less' JSON-based data storage strategy for its audit trails and security logs. "
            "This approach was chosen for its flexibility and ease of integration with JavaScript-based frontends. "
            "The primary data structures are as follows:"
        )
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'File Name'
        hdr[1].text = 'Data Type'
        hdr[2].text = 'Description'
        db_files = [
            ['security_log.json', 'JSON Array', 'Stores threat analysis results, snippets, and timestamps.'],
            ['retraining_log.json', 'JSON Array', 'Stores model versioning, accuracy metrics, and training size.'],
            ['dataset.csv', 'CSV Flat File', 'Primary training data for the Logistic Regression model.'],
            ['user_feedback.csv', 'CSV Flat File', 'Stores user corrections for Active Learning loops.']
        ]
        for f, t, d in db_files:
            row = table.add_row().cells
            row[0].text = f
            row[1].text = t
            row[2].text = d

        doc.add_heading('4.6 Algorithm Detail: Hybrid Orchestration', 2)
        doc.add_paragraph(
            "The core innovation of this project is the orchestration of synchronous and asynchronous tasks. "
            "The algorithm follows a 'Fan-Out' pattern where the search for live news and the LLM analysis are "
            "triggered as concurrent coroutines. This drastically reduces the total turnaround time compared to "
            "sequential execution."
        )
        doc.add_page_break()

    def add_implementation():
        doc.add_heading('CHAPTER 5: IMPLEMENTATION', 1)
        doc.add_paragraph(
            "The implementation phase involved setting up a virtual environment and developing the modular service folder structure. "
            "Each service is self-contained to ensure the 'Single Responsibility Principle'."
        )
        
        doc.add_heading('5.1 ML Service Implementation', 2)
        doc.add_paragraph(
            "The `ml_service.py` uses a Logistic Regression model trained on a curated framing bias dataset. "
            "Text features are extracted using a TF-IDF Vectorizer with n-grams (1,2) to capture local phrases."
        )

        doc.add_heading('5.2 LLM Service Integration', 2)
        doc.add_paragraph(
            "We utilized the GroqCloud SDK to access the LLaMA 3.3 70B model. The integration is fully asynchronous, "
            "allowing the dashboard to remain responsive while the LLM generates deep reasoning reports."
        )
        
        doc.add_heading('5.3 Frontend Dashboard', 2)
        doc.add_paragraph(
            "The dashboard is built with Vanilla JS for performance. It features a dark-themed 'Cyber-Security' aesthetic, "
            "using glassmorphism and dynamic counters to display real-time intelligence."
        )
        doc.add_page_break()

    def add_results_analysis():
        doc.add_heading('CHAPTER 6: RESULTS AND ANALYSIS', 1)
        
        doc.add_heading('6.1 Model Performance', 2)
        doc.add_paragraph("The ML model was evaluated using a 20% hold-out test set. Results are as follows:")
        table = doc.add_table(rows=5, cols=2)
        data = [
            ['Metric', 'Value'],
            ['Accuracy', '83.2%'],
            ['F1-Score', '0.831'],
            ['Precision', '0.832'],
            ['Recall', '0.832']
        ]
        for i, (k, v) in enumerate(data):
            table.cell(i, 0).text = k
            table.cell(i, 1).text = v

        doc.add_heading('6.2 Real-time Analysis Result', 2)
        doc.add_paragraph(
            "During testing, the system correctly identified subtle 'Agenda Framing' in articles related to geopolitical conflicts. "
            "The LLM was able to identify missing context that the live news search retrieved, proving the value of the hybrid approach."
        )
        doc.add_page_break()

    def add_conclusion():
        doc.add_heading('CHAPTER 7: DISCUSSION', 1)
        doc.add_paragraph(
            "The synergy between local ML and cloud LLMs addresses the trade-off between speed and depth. "
            "One challenge faced was the variability in search engine results, which we mitigated by fetching multiple sources "
            "and allowing the LLM to aggregate the consensus."
        )

        doc.add_heading('CHAPTER 8: CONCLUSION & FUTURE SCOPE', 1)
        doc.add_paragraph(
            "This project successfully demonstrates a state-of-the-art approach to misinformation detection. "
            "By combining NLP, Big Data, and AIML, we have created a tool that is both fast and deep."
        )
        doc.add_heading('Future Scope', 2)
        doc.add_paragraph(
            "1. Support for image and video deep-fake detection.\n"
            "2. Integration with browser extensions for real-time browsing protection.\n"
            "3. Collaborative filtering where users can report new bias patterns."
        )
        doc.add_page_break()

    def add_appendix():
        doc.add_heading('REFERENCES', 1)
        doc.add_paragraph("[1] GroqCloud Documentation: API Reference for LLaMA 3.3.")
        doc.add_paragraph("[2] Scikit-Learn: Machine Learning in Python.")
        doc.add_paragraph("[3] Jurafsky, D., & Martin, J. H. (2023). Speech and Language Processing.")
        
        doc.add_heading('APPENDIX', 1)
        doc.add_paragraph("Sample Source Code Snippet (app.py):")
        doc.add_paragraph("app = Flask(__name__)\n@app.route('/api/analyze', methods=['POST'])\ndef analyze():...", style='Quote')
        
    # Build sections
    add_title_page()
    add_abstract()
    add_toc_lists()
    add_introduction()
    add_literature_review()
    add_subject_integration()
    add_system_design()
    add_implementation()
    add_results_analysis()
    add_conclusion()
    add_appendix()

    filename = "Project_Report_Hybrid_AI_Dashboard.docx"
    doc.save(filename)
    print(f"Report generated: {filename}")

if __name__ == "__main__":
    create_report()
